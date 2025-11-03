import PyPDF2
from docx import Document
import logging
import os


logger = logging.getLogger(__name__)


class TextExtractor:
    """Extracts text from PDF and DOCX files"""
    
    def extract_from_pdf(self, file_path: str) -> str:
        """
        Read PDF file and get all text
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            All text from the PDF
        """
        text_parts = []
        
        try:
            logger.info(f"Extracting from PDF: {file_path}")
            
            # Open PDF file
            with open(file_path, 'rb') as file:
                # Create PDF reader object
                pdf_reader = PyPDF2.PdfReader(file)
                
                logger.info(f"PDF has {len(pdf_reader.pages)} pages")
                
                # Go through each page
                for page_num, page in enumerate(pdf_reader.pages):
                    # Extract text from this page
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                        logger.info(f"  Page {page_num + 1}: {len(text)} characters")
            
            total_text = "\n".join(text_parts)
            logger.info(f"✓ Total extracted: {len(total_text)} characters")
            return total_text
        
        except Exception as e:
            logger.error(f"Error reading PDF: {str(e)}")
            raise
    
    def extract_from_docx(self, file_path: str) -> str:
        """
        Read DOCX file and get all text
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            All text from the DOCX
        """
        text_parts = []
        
        try:
            logger.info(f"Extracting from DOCX: {file_path}")
            
            # Open DOCX file
            doc = Document(file_path)
            
            logger.info(f"DOCX has {len(doc.paragraphs)} paragraphs")
            
            # Extract from all paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from all tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text_parts.append(" ".join(row_text))
            
            total_text = "\n".join(text_parts)
            logger.info(f"✓ Total extracted: {len(total_text)} characters")
            return total_text
        
        except Exception as e:
            logger.error(f"Error reading DOCX: {str(e)}")
            raise
    
    def extract_from_file(self, file_path: str) -> str:
        """
        Extract from any supported file
        
        Args:
            file_path: Path to file
            
        Returns:
            Extracted text
        """
        # Check file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Use PDF or DOCX")


# ==================== Test It ====================


if __name__ == "__main__":
    # Test the extractor
    extractor = TextExtractor()
    
    # Test with sample resume
    test_file = "tests/sample_resumes/Shantanu_mamgain_resume.pdf"
    
    if os.path.exists(test_file):
        try:
            text = extractor.extract_from_file(test_file)
            print("\n" + "="*50)
            print("EXTRACTED TEXT")
            print("="*50)
            print(f"Total length: {len(text)} characters")
            print(f"\nFirst 300 characters:")
            print(text[:300])
        except Exception as e:
            print(f"Error: {e}")
    else:
        print(f"Test file not found: {test_file}")
